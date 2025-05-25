# Импортируем модули для создания интерфейса и работы с Word-документами
import tkinter as tk
from tkinter import messagebox
from docx import Document

# Функция для извлечения нужных разделов из исходного .docx по ключевым словам
def extract_sections(doc, keywords):
    result = []              # Список для хранения найденных параграфов
    capture = False          # Флаг захвата текста
    current_keyword = None   # Текущий активный раздел

    # Перебираем все параграфы документа
    for para in doc.paragraphs:
        text = para.text.strip()  # Удаляем пробелы

        # Проверяем начало и конец каждого раздела по ключевым словам
        for keyword in keywords:
            if keyword in text and not text.endswith("конец"):
                capture = True
                current_keyword = keyword
                result.append((para, keyword))  # Сохраняем параграф
                break
            elif keyword in text and text.endswith("конец") and capture and current_keyword == keyword:
                result.append((para, keyword))  # Добавляем завершающий параграф
                capture = False
                current_keyword = None
                break
        else:
            if capture:
                result.append((para, current_keyword))  # Добавляем текст внутри секции

    return result  # Возвращаем список выбранных параграфов

# Функция создания нового документа из выбранных разделов
def generate_doc():
    try:
        src_doc = Document("исходный.docx")  # Загружаем исходный документ
        dst_doc = Document()                 # Создаем новый пустой документ


        # ПРАВИТЬ
        selected_keywords = []              # Список выбранных разделов
        if var_corridor_width.get():
            selected_keywords.append("а1")
        if var_f11_width.get():
            selected_keywords.append("а2")
        if var_f21_zal_width.get():
            selected_keywords.append("а3")
        if var_f21_width_luk_tribun.get():
            selected_keywords.append("а4")
        if var_f21_width_cinema.get():
            selected_keywords.append("а5")
        if var_width_open_tribun.get():
            selected_keywords.append("а6")
        if var_mgn_corridors.get():
            selected_keywords.append("а7")
        if var_corridor_long.get():
            selected_keywords.append("а8")
        if var_f5_corridor_long.get():
            selected_keywords.append("а9")
        if var_f11_width_hosital.get():
            selected_keywords.append("б1")
        if var_f13_corridor_long.get():
            selected_keywords.append("б2")
        if var_open_tribun.get():
            selected_keywords.append("б3")
        if var_f13_corridor_width.get():
            selected_keywords.append("б4")

        if not selected_keywords:
            # Если ничего не выбрано — предупреждение
            messagebox.showwarning("Внимание", "Выберите хотя бы один раздел.")
            return

        # Извлекаем текст из выбранных разделов
        sections = extract_sections(src_doc, selected_keywords)

        # Копируем текст и стили из исходного в новый документ
        for para, _ in sections:
            new_para = dst_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color else None

        dst_doc.save("готовый.docx")  # Сохраняем итоговый файл
        messagebox.showinfo("Готово", "Файл 'готовый.docx' успешно создан.")
    except Exception as e:
        # Показываем сообщение об ошибке, если что-то пошло не так
        messagebox.showerror("Ошибка", str(e))


# === Создание интерфейса ===
root = tk.Tk()                    # Главное окно
root.title("Exert")  # Заголовок окна
root.geometry("600x600")         # Размер окна

# Верхняя рамка с прокруткой
top_frame = tk.Frame(root)
top_frame.pack(fill="both", expand=True)

canvas = tk.Canvas(top_frame)    # Область прокрутки
scrollbar = tk.Scrollbar(top_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)

# Обновляем область прокрутки при изменении содержимого
scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

# Встраиваем прокручиваемый фрейм внутрь канваса
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)  # Размещение канваса
scrollbar.pack(side="right", fill="y")              # Размещение ползунка

# ПРАВИТЬ
# Переменные для чекбоксов
var_corridor_width = tk.BooleanVar()
var_f11_width = tk.BooleanVar()
var_f21_zal_width = tk.BooleanVar()
var_f21_width_luk_tribun = tk.BooleanVar()
var_f21_width_cinema = tk.BooleanVar()
var_width_open_tribun = tk.BooleanVar()
var_mgn_corridors = tk.BooleanVar()
var_corridor_long = tk.BooleanVar()
var_f5_corridor_long = tk.BooleanVar()
var_f11_width_hosital = tk.BooleanVar()
var_f13_corridor_long = tk.BooleanVar()
var_open_tribun = tk.BooleanVar()
var_f13_corridor_width = tk.BooleanVar()

# ПРАВИТЬ
# Подпись и чекбоксы для выбора разделов
tk.Label(scrollable_frame, text="Выберите необходимые параметры:").pack(anchor="w", pady=(5, 5))
tk.Checkbutton(scrollable_frame, text="Ширина коридора для всех зданий", variable=var_corridor_width).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф1.1 Ширина коридора", variable=var_f11_width).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф2.1 Залы ширина проходов", variable=var_f21_zal_width).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф2.1 Залы ширина эвакуационных люков трибун", variable=var_f21_width_luk_tribun).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф2.1 Залы ширина проходов в кинотеатре вместимостью более 100 чел", variable=var_f21_width_cinema).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф2.1 трибуны (залы) на открытом воздухе. Ширина путей эвакуации на трибунах \nв зависимости от числа людей на трибунах", variable=var_width_open_tribun).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Для МГН", variable=var_mgn_corridors).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Коридоры длиной более 60м для всех зданий", variable=var_corridor_long).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф5 Производственные и/или склады. Коридоры длиной более 60м", variable=var_f5_corridor_long).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф1.1 Больницы, специализированные дома престарелых и инвалидов (неквартирные).\nКоридоры длиной более 42м", variable=var_f11_width_hosital).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф1.3 Коридоры длиной более 30м", variable=var_f13_corridor_long).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Трибуны (залы) на открытом воздухе. Ширина путей эвакуации", variable=var_open_tribun).pack(anchor="w")
tk.Checkbutton(scrollable_frame, text="Ф1.3 Ширина коридора в зависимости от длины", variable=var_f13_corridor_width).pack(anchor="w")

# Нижняя рамка — кнопка запуска
bottom_frame = tk.Frame(root)
bottom_frame.pack(fill="x", pady=10)

# Кнопка создания нового документа
tk.Button(bottom_frame, text="Создать готовый.docx", command=generate_doc,
          height=2, font=("Arial", 11, "bold")).pack()

# Запуск графического интерфейса
root.mainloop()

# Напоминание:
# Чтобы встроить иконку в .exe, используй:

# pyinstaller --onefile --noconsole --icon=icon.ico main.py
